from re import S
from altair.vegalite.v4 import api
import streamlit as st
import pandas as pd
from pyairtable import Table
from pyairtable.formulas import match
import io
import os
import base64
import datetime


# Select Date

st.title("Date range")

min_date = datetime.datetime(2020,1,1)
max_date = datetime.date(2022,1,1)

a_date = st.date_input("Pick a date", (min_date, max_date))

# Add Docx library

from docx import Document
from docx.shared import Inches

#Airtable Connection

api_key = st.secrets['airtable_api_key']
base_id = st.secrets['base_id_key']
table_name = st.secrets['table_name_key']


table = Table(api_key, base_id, table_name)

data = table.all()
df = pd.json_normalize(data)


df['fields.Start Date Event'] = pd.to_datetime(df['fields.Start Date']).dt.strftime('%Y-%m-%d')
df['fields.Start Date Time Event'] = pd.to_datetime(df['fields.Start Date']).dt.strftime('%H:%M:%S')

df_columns = df[['fields.Event Title', 'fields.Agenda & Description & Blurb', 'fields.POC', 'fields.Start Date Event']]

df_filtered = df_columns[(df_columns['fields.Start Date Event'] > a_date[0].strftime('%Y-%m-%d')) & (df_columns['fields.Start Date Event'] < a_date[1].strftime('%Y-%m-%d'))]

st.dataframe(df_filtered)

# st.dataframe(df[['fields.Event Title', 'fields.Start Date Event', 'fields.Start Date Time Event', 'fields.Room/Platform', 'fields.POC']][df['fields.Start Date Event'] == a_date[0]])

# Docs Creation

document = Document()

def CreateDocs():

    document.add_heading('For RSO', 0)

    for row in df_filtered.itertuples():
        document.add_paragraph(row[3])

    document.save('For RSO.docx')

st.button(label = 'Generate DOCS', on_click=CreateDocs())

import base64
import os
import json
import pickle
import uuid
import re

import streamlit as st
import pandas as pd


def download_button(object_to_download, download_filename, button_text, pickle_it=False):
    """
    Generates a link to download the given object_to_download.
    Params:
    ------
    object_to_download:  The object to be downloaded.
    download_filename (str): filename and extension of file. e.g. mydata.csv,
    some_txt_output.txt download_link_text (str): Text to display for download
    link.
    button_text (str): Text to display on download button (e.g. 'click here to download file')
    pickle_it (bool): If True, pickle file.
    Returns:
    -------
    (str): the anchor tag to download object_to_download
    Examples:
    --------
    download_link(your_df, 'YOUR_DF.csv', 'Click to download data!')
    download_link(your_str, 'YOUR_STRING.txt', 'Click to download text!')
    """
    if pickle_it:
        try:
            object_to_download = pickle.dumps(object_to_download)
        except pickle.PicklingError as e:
            st.write(e)
            return None

    else:
        if isinstance(object_to_download, bytes):
            pass

        elif isinstance(object_to_download, pd.DataFrame):
            object_to_download = object_to_download.to_csv(index=False)

        # Try JSON encode for everything else
        else:
            object_to_download = json.dumps(object_to_download)

    try:
        # some strings <-> bytes conversions necessary here
        b64 = base64.b64encode(object_to_download.encode()).decode()

    except AttributeError as e:
        b64 = base64.b64encode(object_to_download).decode()

    button_uuid = str(uuid.uuid4()).replace('-', '')
    button_id = re.sub('\d+', '', button_uuid)

    custom_css = f""" 
        <style>
            #{button_id} {{
                background-color: rgb(255, 255, 255);
                color: rgb(38, 39, 48);
                padding: 0.25em 0.38em;
                position: relative;
                text-decoration: none;
                border-radius: 4px;
                border-width: 1px;
                border-style: solid;
                border-color: rgb(230, 234, 241);
                border-image: initial;
            }} 
            #{button_id}:hover {{
                border-color: rgb(246, 51, 102);
                color: rgb(246, 51, 102);
            }}
            #{button_id}:active {{
                box-shadow: none;
                background-color: rgb(246, 51, 102);
                color: white;
                }}
        </style> """

    dl_link = custom_css + f'<a download="{download_filename}" id="{button_id}" href="data:file/txt;base64,{b64}">{button_text}</a><br></br>'

    return dl_link


# def file_selector(folder_path='.'):
#     filenames = os.listdir(folder_path)
#     selected_filename = st.selectbox('Select a file', filenames)
#     return os.path.join(folder_path, selected_filename)


if __name__ == '__main__':
    
    # Upload file for testing
        # folder_path = st.text_input('Enter directory: deafult .', '.')
        # filename = file_selector(folder_path=folder_path)
        filename = './For RSO.docx'

        # Load selected file
        with open(filename, 'rb') as f:
            s = f.read()

        download_button_str = download_button(s, filename, f'Click here to download {filename}')
        st.markdown(download_button_str, unsafe_allow_html=True)


        
